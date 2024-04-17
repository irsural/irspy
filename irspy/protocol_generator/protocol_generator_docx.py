from pathlib import Path
from typing import Iterable
from zipfile import BadZipFile
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from irspy.protocol_generator.data_table import DataTable, DataCell
from irspy.protocol_generator.protocol_generator import ProtocolGenerator, ProtocolConfig
from irspy.protocol_generator.utils import handle_exceptions


class DocxProtocolGenerator(ProtocolGenerator):
    def __init__(self) -> None:
        pass

    @handle_exceptions(BadZipFile, PermissionError, PackageNotFoundError)
    def generate(self, file: Path, config: ProtocolConfig) -> None:
        document = Document(file)

        for tag, paragraph in self.__find_tagged_paragraphs(document, config.tag_map.keys()):
            value = config.tag_map[tag]
            if isinstance(value, DataTable):
                docx_table = self.__make_table(document, value)
                self.__insert_table(docx_table, paragraph)
                self.__remove_paragraph(paragraph)
            else:
                paragraph.text = paragraph.text.replace(tag, str(value))

        document.save(file)

    def __set_repeat_table_header(self, row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def __find_tagged_paragraphs(self, document, tags: Iterable[str]):
        def find_paragraphs(element):
            for p in element.paragraphs:
                for tag in tags:
                    if tag in p.text:
                        yield tag, p

        yield from find_paragraphs(document)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from find_paragraphs(cell)

    def __make_table(self, document, data_table: DataTable):
        table = document.add_table(rows=data_table.row_count, cols=data_table.column_count)
        table_cells = table._cells
        if data_table.add_border:
            try:
                table.style = 'Table Grid'
            except KeyError as ex:
                print(str(ex))
        for data_cell in data_table.cells:
            start_cell = table_cells[data_cell.row * data_table.column_count + data_cell.column]
            start_cell.text = data_cell.value
            if data_cell.row_span > 1 or data_cell.column_span > 1:
                end_cell = table_cells[(
                                                   data_cell.row + data_cell.row_span - 1) * data_table.column_count + data_cell.column + data_cell.column_span - 1]
                start_cell.merge(end_cell)

        for i in range(data_table.header_rows):
            self.__set_repeat_table_header(table.rows[i])

        return table

    def __insert_table(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def __remove_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None